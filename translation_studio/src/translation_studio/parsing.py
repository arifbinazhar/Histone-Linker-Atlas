from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import BinaryIO

from .models import Segment


SEGMENT_SPLIT_RE = re.compile(r"(?<=[.!?。！？])\s+(?=[A-Z0-9\"'“‘])")


def parse_uploaded_file(file_obj: BinaryIO, filename: str) -> list[Segment]:
    suffix = Path(filename).suffix.lower()
    with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_obj.read())
        tmp_path = Path(tmp.name)

    try:
        if suffix == ".docx":
            return parse_docx(tmp_path)
        if suffix == ".pdf":
            return parse_pdf(tmp_path)
        if suffix in {".txt", ".md"}:
            return parse_text(tmp_path.read_text(encoding="utf-8", errors="ignore"))
        raise ValueError("Unsupported file type. Upload a PDF, DOCX, or TXT file.")
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_docx(path: Path) -> list[Segment]:
    try:
        from docx import Document
    except ImportError:
        return parse_docx_with_stdlib(path)

    document = Document(path)
    blocks: list[tuple[str, str, str]] = []
    current_section = ""

    for paragraph in document.paragraphs:
        text = normalize_extracted_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        block_type = "heading" if "heading" in style_name or style_name.startswith("title") else "paragraph"
        if block_type == "heading":
            current_section = text
        blocks.append((text, block_type, current_section))

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = normalize_extracted_text(cell.text)
                if text:
                    section = current_section or f"Table {table_index}"
                    block = f"table cell {table_index}.{row_index}.{cell_index}"
                    blocks.append((text, block, section))

    return blocks_to_segments(blocks)


def parse_pdf(path: Path) -> list[Segment]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return parse_pdf_best_effort(path)

    reader = PdfReader(str(path))
    blocks: list[tuple[str, str, str]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for line in text.splitlines():
            normalized = normalize_extracted_text(line)
            if normalized:
                block_type = "heading" if looks_like_heading(normalized) else "paragraph"
                blocks.append((normalized, f"page {page_index} {block_type}", f"Page {page_index}"))
    return blocks_to_segments(blocks)


def parse_text(text: str) -> list[Segment]:
    blocks = [(normalize_extracted_text(line), "paragraph", "") for line in text.splitlines()]
    return blocks_to_segments([block for block in blocks if block[0]])


def parse_docx_with_stdlib(path: Path) -> list[Segment]:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    blocks: list[tuple[str, str, str]] = []
    current_section = ""
    with zipfile.ZipFile(path) as archive:
        xml_payload = archive.read("word/document.xml")
    root = ET.fromstring(xml_payload)
    body = root.find("w:body", namespace)
    if body is None:
        return []

    for child in body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = paragraph_text(child, namespace)
            if not text:
                continue
            block_type = "heading" if looks_like_heading(text) else "paragraph"
            if block_type == "heading":
                current_section = text
            blocks.append((text, block_type, current_section))
        elif tag == "tbl":
            for cell in child.findall(".//w:tc", namespace):
                text = normalize_extracted_text(" ".join(paragraph_text(p, namespace) for p in cell.findall(".//w:p", namespace)))
                if text:
                    blocks.append((text, "table cell", current_section))
    return blocks_to_segments(blocks)


def paragraph_text(paragraph: ET.Element, namespace: dict[str, str]) -> str:
    pieces = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
    return normalize_extracted_text("".join(pieces))


def parse_pdf_best_effort(path: Path) -> list[Segment]:
    raw = path.read_bytes()
    printable = re.sub(rb"[^ -~\r\n\t]", b" ", raw).decode("latin-1", errors="ignore")
    candidates = re.findall(r"\(([^()]{3,500})\)\s*Tj", printable)
    if not candidates:
        candidates = re.findall(r"\[([^\]]{3,1200})\]\s*TJ", printable)
    text = "\n".join(normalize_extracted_text(item.replace("\\)", ")").replace("\\(", "(")) for item in candidates)
    if not text.strip():
        raise RuntimeError("PDF parsing needs pypdf for this file. Install optional dependencies with: pip install -r requirements.txt")
    return parse_text(text)


def blocks_to_segments(blocks: list[tuple[str, str, str]]) -> list[Segment]:
    segments: list[Segment] = []
    for text, block_type, section in blocks:
        parts = split_segment_text(text) if block_type != "heading" else [text]
        for part in parts:
            if len(part.strip()) >= 2:
                segments.append(
                    Segment(
                        text=part.strip(),
                        block_type=block_type,
                        section=section,
                        index=len(segments) + 1,
                    )
                )
    return segments


def split_segment_text(text: str) -> list[str]:
    if len(text) < 220:
        return [text]
    return [part.strip() for part in SEGMENT_SPLIT_RE.split(text) if part.strip()]


def normalize_extracted_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def looks_like_heading(text: str) -> bool:
    words = text.split()
    if not words or len(words) > 12:
        return False
    return text.isupper() or text.istitle() or text.endswith(":")
